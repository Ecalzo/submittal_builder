'''
Builds a cover sheet for S.P. Contractors from scratch
'''
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import win32com.client as client
import os
import datetime

def buildCoverPage(submittal_no='', project_name="8 East 3rd St.", product_name = ''):
    #TODO: receive inputs for submittal no. etc..
    # gather today's date and format the string
    todays_date = datetime.datetime.now().strftime("%m-%d-%Y")
    # Initialize document object
    document = Document()
    heading = document.add_heading("S.P. Contractors", 0)

    address = document.add_paragraph()
    address_text = address.add_run("Queens, NY")

    #whitestpace
    document.add_paragraph("\n\n\n\n\n\n\n")

    doc_body_text = "Project: " + project_name + "\n" + "Submittal No.:"  + submittal_no + "\n" + "Date: " + str(todays_date) + '\n'

    body = document.add_paragraph()

    # Align text center
    body_format = body.paragraph_format
    body_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # seperate the text from the paragraph structure so we can apply font changes
    body_text = body.add_run(doc_body_text)
    font = body_text.font
    font.name = 'Calibri'
    font.size = Pt(24)

    # build file name and save the file, print success statement
    file_name = submittal_no + ' ' + product_name + ' ' + project_name + '.docx'
    document.save(file_name)
    print("Successfully output ", file_name)
